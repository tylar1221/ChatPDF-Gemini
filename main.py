import streamlit as st
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain_community.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import joblib
from docx import Document
from io import BytesIO
from fpdf import FPDF

# Load environment variables
load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# File to store chat history
CHAT_HISTORY_FILE = "chat_history.pkl"

# Initialize or load chat history
if os.path.exists(CHAT_HISTORY_FILE):
    chat_history = joblib.load(CHAT_HISTORY_FILE)
else:
    chat_history = []


def get_pdf_text(pdf_docs):
    text = ""
    for pdf in pdf_docs:
        pdf_reader = PdfReader(pdf)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text


def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=1000)
    chunks = text_splitter.split_text(text)
    return chunks


def get_vector_store(text_chunks):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")


def get_conversational_chain():
    prompt_template = """
    Answer the question as detailed as possible from the provided context in simple words, make sure to provide all the details, if the answer is not in
    provided context just say, "answer is not available in the context", don't provide the wrong answer\n\n
    Context:\n {context}?\n
    Question: \n{question}\n

    Answer:
    """

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.5)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

    return chain


def user_input(user_question):
    global chat_history

    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)
    docs = new_db.similarity_search(user_question)

    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)

    # Save chat history
    chat_entry = {"question": user_question, "answer": response["output_text"]}
    chat_history.append(chat_entry)
    joblib.dump(chat_history, CHAT_HISTORY_FILE)

    st.write("Reply: ", response["output_text"])


def download_chat_history(format_choice):
    if format_choice == "Word":
        doc = Document()
        doc.add_heading("Chat History", level=1)
        for entry in chat_history:
            doc.add_paragraph(f"Q: {entry['question']}")
            doc.add_paragraph(f"A: {entry['answer']}")
            doc.add_paragraph("---")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer, "chat_history.docx"

    elif format_choice == "PDF":
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Chat History", ln=True, align="C")
        for entry in chat_history:
            pdf.multi_cell(0, 10, f"Q: {entry['question']}")
            pdf.multi_cell(0, 10, f"A: {entry['answer']}")
            pdf.cell(0, 10, "---", ln=True)
        buffer = BytesIO()

        pdf_output = pdf.output(dest='S').encode('latin1')
        buffer.write(pdf_output)
        buffer.seek(0)

        return buffer, "chat_history.pdf"


def main():
    st.set_page_config(page_title="Chat PDF", layout="wide")
    st.header("Chat with PDF using Gemini💁")

    user_question = st.text_input("Ask a Question from the PDF Files")

    if user_question:
        user_input(user_question)

    with st.sidebar:
        st.title("Menu:")
        pdf_docs = st.file_uploader("Upload your PDF Files and Click on the Submit & Process Button", accept_multiple_files=True)
        if st.button("Submit & Process"):
            with st.spinner("Processing..."):
                raw_text = get_pdf_text(pdf_docs)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Processing Completed!")

        # Display chat history
        if chat_history:
            st.subheader("Chat History")
            for entry in chat_history:
                st.write(f"**Q:** {entry['question']}")
                st.write(f"**A:** {entry['answer']}")
                st.write("---")

            # Download options
            format_choice = st.selectbox("Select format to download chat history:", ["Word", "PDF"])
            if st.button("Download"):
                buffer, filename = download_chat_history(format_choice)
                st.download_button(label="Download Chat History", data=buffer, file_name=filename)

        # Clear history button
        if st.button("Clear History"):
            if os.path.exists(CHAT_HISTORY_FILE):
                os.remove(CHAT_HISTORY_FILE)
                chat_history.clear()
                st.success("Chat history cleared successfully!")
            else:
                st.warning("No chat history to clear.")


if __name__ == "__main__":
    main()
